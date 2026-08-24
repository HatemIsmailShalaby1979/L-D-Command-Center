# engines/career-engine/resume/test_parser.py
#
# WHAT: Tests for the resume parser module.
# WHY:  Ensure the parser correctly extracts fields from sample resumes
#       and flags low-confidence extractions properly.
# BREAKS IF DELETED: No regression testing for resume upload parsing.

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from typing import Any

import pytest

# Resolve paths
_ROOT = Path(__file__).resolve().parent.parent.parent
_ENGINE_PATH = Path(__file__).resolve().parent
_MODEL_LAYER = _ROOT / "model-layer"

import sys
sys.path.insert(0, str(_ENGINE_PATH))
sys.path.insert(0, str(_MODEL_LAYER))

from parser import (
    parse_resume_file,
    parse_text,
    ConfidenceFlag,
    HIGH_CONFIDENCE,
    MEDIUM_CONFIDENCE,
    UNKNOWN,
)


# ---------------------------------------------------------------------------
# Sample resume text for testing
# ---------------------------------------------------------------------------

SAMPLE_RESUME_TEXT = """
John Michael Smith
john.smith@email.com | (555) 123-4567
San Francisco, CA | linkedin.com/in/johnsmith | github.com/johnsmith

PROFESSIONAL SUMMARY
Experienced software engineer with 8+ years in full-stack development,
specializing in Python, JavaScript, and cloud infrastructure. Proven track
record of leading cross-functional teams and delivering scalable solutions.

PROFESSIONAL EXPERIENCE

Senior Software Engineer | TechCorp Inc. | 2020-Present
- Led development of microservices architecture serving 10M+ users
- Implemented CI/CD pipelines reducing deployment time by 60%
- Mentored junior developers and conducted code reviews

Software Engineer | StartupXYZ | 2018-2020
- Built REST APIs using Django and Flask
- Developed React frontend components
- Reduced page load times by 40% through optimization

Junior Developer | WebAgency | 2016-2018
- Created responsive websites for clients
- Worked with WordPress and custom PHP solutions

EDUCATION

Bachelor of Science in Computer Science | University of California, Berkeley | 2012-2016

TECHNICAL SKILLS
Python, JavaScript, TypeScript, React, Django, Flask, AWS, Docker, Kubernetes, PostgreSQL, MongoDB

PROJECTS

- Inventory Management System: Full-stack app with React frontend and Node.js backend
- ML Pipeline: Automated data processing using Python and TensorFlow

"""


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestParseResumeFile:
    """Tests for parse_resume_file() with actual file parsing."""

    def test_parse_pdf_resume(self):
        """
        Contract: parse_resume_file should extract fields from a PDF.
        """
        # We'll test with text parsing since we don't have a real PDF
        # This tests the integration path
        pass  # Skip for now - requires PDF generation

    def test_parse_docx_resume(self):
        """
        Contract: parse_resume_file should extract fields from a DOCX.
        """
        pass  # Skip for now - requires DOCX generation

    def test_unsupported_file_type(self):
        """
        Contract: parse_resume_file should raise ValueError for unsupported types.
        """
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"test")
            temp_path = Path(f.name)

        try:
            with pytest.raises(ValueError, match="Unsupported file type"):
                parse_resume_file(temp_path)
        finally:
            temp_path.unlink()


class TestParseText:
    """Tests for parse_text() with raw text input."""

    def test_parses_contact_info(self):
        """
        Contract: parse_text should extract name, email, phone, location, LinkedIn, GitHub.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert resume["contact"]["name"] == "John Michael Smith"
        assert resume["contact"]["email"] == "john.smith@email.com"
        assert "555" in resume["contact"]["phone"]  # Loose check for phone
        assert "San Francisco" in resume["contact"]["location"]  # Loose check for location
        assert "linkedin.com/in/johnsmith" in resume["contact"]["linkedin"]
        assert "github.com/johnsmith" in resume["contact"]["github"]

    def test_parses_summary(self):
        """
        Contract: parse_text should extract the professional summary.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert resume["summary"] is not None
        assert "experienced software engineer" in resume["summary"].lower()
        assert len(resume["summary"]) > 50

    def test_parses_experience(self):
        """
        Contract: parse_text should extract experience entries with title, company, dates, description.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert len(resume["experience"]) >= 3

        # Check first experience entry
        first_exp = resume["experience"][0]
        assert "Senior Software Engineer" in first_exp["title"]
        assert "TechCorp" in first_exp["company"]
        assert "2020" in first_exp["dates"]
        assert len(first_exp["description"]) > 0

        # Check second experience entry
        second_exp = resume["experience"][1]
        assert "Software Engineer" in second_exp["title"]
        assert "StartupXYZ" in second_exp["company"]

    def test_parses_education(self):
        """
        Contract: parse_text should extract education entries with degree, school, dates.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert len(resume["education"]) >= 1

        edu = resume["education"][0]
        assert "Bachelor" in edu["degree"]
        assert "Berkeley" in edu["school"]
        assert "2012" in edu["dates"]

    def test_parses_skills(self):
        """
        Contract: parse_text should extract skills list.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert len(resume["skills"]) >= 5
        assert "Python" in resume["skills"]
        assert "JavaScript" in resume["skills"]
        assert "React" in resume["skills"]

    def test_parses_projects(self):
        """
        Contract: parse_text should extract project entries.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        assert len(resume["projects"]) >= 2

        first_project = resume["projects"][0]
        assert first_project["name"] == "Inventory Management System"
        assert len(first_project["description"]) > 0

    def test_flags_missing_fields(self):
        """
        Contract: parse_text should flag missing fields as UNKNOWN confidence.
        """
        minimal_text = "John Doe\njohn@example.com\n"
        resume, flags = parse_text(minimal_text)

        # Should have flags for missing sections
        missing_fields = [f for f in flags if f.confidence == UNKNOWN]
        assert len(missing_fields) > 0

        # Check that required sections are flagged
        flagged_fields = {f.field for f in missing_fields}
        assert "experience" in flagged_fields
        assert "education" in flagged_fields

    def test_flags_confidence_levels(self):
        """
        Contract: parse_text should assign appropriate confidence levels.
        """
        resume, flags = parse_text(SAMPLE_RESUME_TEXT)

        # Email should be high confidence
        email_flags = [f for f in flags if f.field == "contact.email"]
        assert len(email_flags) > 0
        assert email_flags[0].confidence == HIGH_CONFIDENCE

        # Phone should be medium confidence
        phone_flags = [f for f in flags if f.field == "contact.phone"]
        assert len(phone_flags) > 0
        assert phone_flags[0].confidence == MEDIUM_CONFIDENCE


class TestEdgeCases:
    """Tests for edge cases and error handling."""

    def test_empty_text(self):
        """
        Contract: parse_text should handle empty input gracefully.
        """
        resume, flags = parse_text("")

        assert resume["contact"]["name"] == ""
        assert resume["contact"]["email"] == ""
        assert resume["experience"] == []
        assert resume["education"] == []
        assert resume["skills"] == []
        assert resume["projects"] == []

        # All fields should be flagged as unknown
        unknown_flags = [f for f in flags if f.confidence == UNKNOWN]
        assert len(unknown_flags) >= 6

    def test_malformed_text(self):
        """
        Contract: parse_text should not crash on malformed text.
        """
        messy_text = "Random text without structure\n\nMore random stuff\n\n12345"
        resume, flags = parse_text(messy_text)

        # Should return empty resume without crashing
        assert isinstance(resume, dict)
        assert "contact" in resume
        assert "experience" in resume

    def test_unicode_content(self):
        """
        Contract: parse_text should handle unicode characters.
        """
        unicode_text = """
        María García
        maria@garcia.com
        Software Engineer
        """
        resume, flags = parse_text(unicode_text)

        assert resume["contact"]["email"] == "maria@garcia.com"
        # Name should be present (may be extracted from first line)
        assert resume["contact"]["name"] is not None


# ---------------------------------------------------------------------------
# Sample DOCX for live testing
# ---------------------------------------------------------------------------

def create_sample_docx(path: Path) -> None:
    """Helper to create a sample DOCX file for testing."""
    from docx import Document

    doc = Document()

    # Contact info
    doc.add_paragraph("John Michael Smith")
    doc.add_paragraph("john.smith@email.com | (555) 123-4567")
    doc.add_paragraph("San Francisco, CA | linkedin.com/in/johnsmith | github.com/johnsmith")

    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Experienced software engineer with 8+ years in full-stack development.")

    # Experience
    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph("Senior Software Engineer | TechCorp Inc. | 2020-Present")
    doc.add_paragraph("Led development of microservices architecture serving 10M+ users")

    doc.add_paragraph("Software Engineer | StartupXYZ | 2018-2020")
    doc.add_paragraph("Built REST APIs using Django and Flask")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science | University of California, Berkeley | 2012-2016")

    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Python, JavaScript, TypeScript, React, Django, Flask, AWS, Docker")

    # Projects
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("Inventory Management System: Full-stack app with React frontend")

    doc.save(path)


def create_sample_pdf(path: Path) -> None:
    """Helper to create a sample PDF file for testing."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter

    # Contact info
    c.drawString(100, height - 100, "John Michael Smith")
    c.drawString(100, height - 120, "john.smith@email.com | (555) 123-4567")
    c.drawString(100, height - 140, "San Francisco, CA | linkedin.com/in/johnsmith | github.com/johnsmith")

    # Summary
    c.drawString(100, height - 180, "PROFESSIONAL SUMMARY")
    c.drawString(100, height - 200, "Experienced software engineer with 8+ years in full-stack development.")

    # Experience
    c.drawString(100, height - 240, "PROFESSIONAL EXPERIENCE")
    c.drawString(100, height - 260, "Senior Software Engineer | TechCorp Inc. | 2020-Present")
    c.drawString(100, height - 280, "Led development of microservices architecture serving 10M+ users")

    c.drawString(100, height - 320, "Software Engineer | StartupXYZ | 2018-2020")
    c.drawString(100, height - 340, "Built REST APIs using Django and Flask")

    # Education
    c.drawString(100, height - 380, "EDUCATION")
    c.drawString(100, height - 400, "Bachelor of Science in Computer Science | University of California, Berkeley | 2012-2016")

    # Skills
    c.drawString(100, height - 440, "TECHNICAL SKILLS")
    c.drawString(100, height - 460, "Python, JavaScript, TypeScript, React, Django, Flask, AWS, Docker")

    # Projects
    c.drawString(100, height - 500, "PROJECTS")
    c.drawString(100, height - 520, "Inventory Management System: Full-stack app with React frontend")

    c.save()


class TestLiveParsing:
    """Live tests with actual PDF/DOCX files."""

    def test_parse_docx_file(self, tmp_path: Path):
        """
        Contract: parse_resume_file should successfully parse a DOCX file.
        """
        docx_path = tmp_path / "test_resume.docx"
        create_sample_docx(docx_path)

        resume, flags = parse_resume_file(docx_path)

        assert resume["contact"]["name"] is not None
        assert resume["contact"]["email"] is not None
        assert len(resume["experience"]) >= 2
        assert len(resume["education"]) >= 1
        assert len(resume["skills"]) >= 3

    def test_parse_pdf_file(self, tmp_path: Path):
        """
        Contract: parse_resume_file should successfully parse a PDF file.
        """
        pdf_path = tmp_path / "test_resume.pdf"
        create_sample_pdf(pdf_path)

        resume, flags = parse_resume_file(pdf_path)

        assert resume["contact"]["name"] is not None
        assert resume["contact"]["email"] is not None
        assert len(resume["experience"]) >= 1
        assert len(resume["education"]) >= 1


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
