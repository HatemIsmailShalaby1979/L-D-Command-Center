# engines/export-engine/docx_format.py
#
# WHAT: Deterministic DOCX renderer and writer (Resume).
# WHY:  P4.1 split of the former 1004-line export god-module (CONSTITUTION
#       §4): one file per concern so layout bugs localize to one adapter.
# BREAKS IF DELETED: DOCX exports disappear.

from __future__ import annotations

import logging
from datetime import datetime
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)

_EPOCH = datetime(2000, 1, 1)

def export_resume_to_docx(resume: dict[str, Any]) -> bytes:
    """
    Contract: render a Resume dict to a DOCX document (bytes).
    Uses python-docx to generate a professional resume document.

    Args:
        resume: a validated Resume dict.

    Returns:
        DOCX bytes ready to write to disk.

    Raises:
        ValueError: if the resume is missing required fields.
        ImportError: if python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError(
            "docx export requires python-docx: pip install python-docx"
        ) from exc

    contact = resume.get("contact", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    projects = resume.get("projects", [])

    if not contact.get("name"):
        raise ValueError("Resume must have a contact.name field")
    if not experience:
        raise ValueError("Resume must contain at least one experience entry")

    doc = Document()
    props = doc.core_properties
    props.author = "L&D Command Center"
    props.last_modified_by = "L&D Command Center"
    props.created = _EPOCH   # byte-stability
    props.modified = _EPOCH

    # Name
    name = contact.get("name", "Name")
    title_para = doc.add_heading(name, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact_parts = []
    for key in ["email", "phone", "location", "linkedin", "github"]:
        val = contact.get(key, "")
        if val:
            contact_parts.append(val)
    if contact_parts:
        contact_para = doc.add_paragraph("  |  ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    if summary:
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(summary)

    # Experience
    if experience:
        doc.add_heading("EXPERIENCE", level=1)
        for exp in experience:
            title = exp.get("title", "Title")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "Dates")
            description = exp.get("description", "")

            exp_para = doc.add_paragraph()
            exp_para.add_run(f"{title} | {company} | {dates}").bold = True

            for line in description.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(f"- {line}", style="List Bullet")

    # Education
    if education:
        doc.add_heading("EDUCATION", level=1)
        for edu in education:
            degree = edu.get("degree", "Degree")
            school = edu.get("school", "School")
            dates = edu.get("dates", "Dates")
            doc.add_paragraph(f"{degree} | {school} | {dates}")

    # Skills
    if skills:
        doc.add_heading("SKILLS", level=1)
        doc.add_paragraph(", ".join(skills))

    # Projects
    if projects:
        doc.add_heading("PROJECTS", level=1)
        for proj in projects:
            proj_name = proj.get("name", "Project")
            desc = proj.get("description", "")
            tech = proj.get("tech", [])

            proj_para = doc.add_paragraph()
            proj_para.add_run(proj_name).bold = True

            if desc:
                doc.add_paragraph(desc)
            if tech:
                doc.add_paragraph(f"Tech: {', '.join(tech)}")

    # Serialize to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_resume_to_docx_file(resume: dict[str, Any], filepath: str) -> str:
    """
    Contract: write resume DOCX export to a file.

    Args:
        resume: a validated Resume dict.
        filepath: path to write the .docx file.

    Returns:
        The filepath that was written.
    """
    docx_bytes = export_resume_to_docx(resume)
    with open(filepath, "wb") as f:
        f.write(docx_bytes)
    logger.info("Wrote resume DOCX export to %s (%d bytes)", filepath, len(docx_bytes))
    return filepath


# ---------------------------------------------------------------------------
# Audio exports
# ---------------------------------------------------------------------------
