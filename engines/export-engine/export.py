# engines/export-engine/export.py
#
# WHAT: Export module for the export-engine — converts validated Journey
#       or Resume dicts into plain-text, PDF, and DOCX output formats.
# WHY:  Consumers need downloadable formats beyond HTML (journeys) or
#       editable documents (resumes). Both are deterministic (no model
#       calls), so they can be unit-tested and cached.
# BREAKS IF DELETED: The export-engine loses its primary outputs.
#       Users cannot share or download content outside the browser.

from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Type detection
# ---------------------------------------------------------------------------

def _detect_type(data: dict[str, Any]) -> str:
    """
    Contract: detect whether the data dict is a Journey, Resume, or audio type.

    Detection heuristic:
      - 'topic' + 'cards' keys → 'journey'
      - 'contact' + 'experience' keys → 'resume'
      - 'target_language' + 'known_language' keys → 'bilingual'
      - 'topic' + 'target_language' + 'segments' keys → 'immersion'
      - 'topic' + 'host_name' keys (no 'cards') → 'podcast'
      - 'text' key (and not any of the above) → 'narration'
      - otherwise → 'unknown'

    Returns:
        'journey', 'resume', 'bilingual', 'immersion', 'podcast',
        'narration', or 'unknown'
    """
    if "topic" in data and "cards" in data:
        return "journey"
    if "contact" in data and "experience" in data:
        return "resume"
    if "target_language" in data and "known_language" in data:
        return "bilingual"
    if "topic" in data and "target_language" in data and "segments" in data:
        return "immersion"
    if "topic" in data and "host_name" in data:
        return "podcast"
    if "text" in data:
        return "narration"
    return "unknown"


# ---------------------------------------------------------------------------
# Journey exports
# ---------------------------------------------------------------------------

def export_to_plain_text(journey: dict[str, Any]) -> str:
    """
    Contract: render a Journey dict to a plain-text string.

    Format:
      TITLE (LEVEL)
      ==============

      Card 1: <title>
      <content>

      Quiz: <question>
      A) <option 0>
      B) <option 1>
      C) <option 2>
      D) <option 3>
      Correct: <correct_option>
      Explanation: <explanation>

      ---

      [repeated for each card]

    Args:
        journey: a validated Journey dict with keys 'topic',
                 'level', and 'cards'.

    Returns:
        A plain-text string.

    Raises:
        ValueError: if the journey is missing required fields.
    """
    topic = journey.get("topic")
    level = journey.get("level", "beginner")
    cards = journey.get("cards", [])

    if not topic:
        raise ValueError("Journey must have a 'topic' field")
    if not cards:
        raise ValueError("Journey must contain at least one card")

    lines = [
        f"{topic.upper()} ({level.upper()})",
        "=" * len(topic),
        "",
    ]

    for i, card in enumerate(cards, start=1):
        card_id = card.get("id", str(i))
        title = card.get("title", "Untitled Card")
        content = card.get("content", "")
        question = card.get("question", "")
        options = card.get("options", [])
        correct_option = card.get("correct_option", "")
        explanation = card.get("explanation", "")

        lines.append(f"Card {i}: {title}")
        lines.append(content)
        lines.append("")
        lines.append(f"Quiz: {question}")

        for j, opt in enumerate(options):
            label = chr(65 + j)  # A, B, C, D...
            lines.append(f"  {label}) {opt}")

        lines.append(f"  Correct: {correct_option}")
        lines.append(f"  Explanation: {explanation}")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def export_to_pdf(journey: dict[str, Any]) -> bytes:
    """
    Contract: render a Journey dict to a PDF document (bytes).
    Uses fpdf2 to generate a clean, multi-page PDF.

    Args:
        journey: a validated Journey dict.

    Returns:
        PDF bytes ready to write to disk or serve.

    Raises:
        ValueError: if the journey is missing required fields.
        ImportError: if fpdf2 is not installed.
    """
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ImportError(
            "pdfexport requires fpdf2: pip install fpdf2"
        ) from exc

    topic = journey.get("topic")
    level = journey.get("level", "beginner")
    cards = journey.get("cards", [])

    if not topic:
        raise ValueError("Journey must have a 'topic' field")
    if not cards:
        raise ValueError("Journey must contain at least one card")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, topic, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Level badge
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(67, 97, 238)  # accent color
    pdf.cell(0, 8, f"Level: {level.title()}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    for i, card in enumerate(cards, start=1):
        # If not first card, add a page break if needed (fpdf handles automatically)
        if i > 1:
            pdf.add_page()

        # Card title
        pdf.set_font("Helvetica", "B", 14)
        card_id = card.get("id", str(i))
        title = card.get("title", f"Card {i}")
        pdf.cell(0, 10, f"Card {i}: {title}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # Content
        pdf.set_font("Helvetica", "", 11)
        content = card.get("content", "")
        pdf.multi_cell(0, 6, content)
        pdf.ln(4)

        # Quiz question
        pdf.set_font("Helvetica", "B", 11)
        question = card.get("question", "")
        pdf.cell(0, 6, f"Quiz: {question}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # Options
        pdf.set_font("Helvetica", "", 11)
        options = card.get("options", [])
        correct_option = card.get("correct_option", "")
        for j, opt in enumerate(options):
            label = f"{chr(65 + j)})"
            pdf.cell(8, 6, label)
            pdf.cell(0, 6, opt, new_x="END", new_y="NEXT")
        pdf.ln(3)

        # Correct answer and explanation
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(46, 196, 182)  # success color
        pdf.cell(0, 6, f"Correct: {correct_option}")
        pdf.ln(4)
        pdf.set_text_color(108, 117, 125)  # muted color
        explanation = card.get("explanation", "")
        pdf.multi_cell(0, 5, f"Explanation: {explanation}")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(6)

    # Serialize to bytes
    buf = BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def export_to_text_file(journey: dict[str, Any], filepath: str) -> str:
    """
    Contract: write plain-text export to a file.

    Args:
        journey: a validated Journey dict.
        filepath: path to write the .txt file.

    Returns:
        The filepath that was written.
    """
    text = export_to_plain_text(journey)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)
    logger.info("Wrote plain-text export to %s", filepath)
    return filepath


def export_to_pdf_file(journey: dict[str, Any], filepath: str) -> str:
    """
    Contract: write PDF export to a file.

    Args:
        journey: a validated Journey dict.
        filepath: path to write the .pdf file.

    Returns:
        The filepath that was written.
    """
    pdf_bytes = export_to_pdf(journey)
    with open(filepath, "wb") as f:
        f.write(pdf_bytes)
    logger.info("Wrote PDF export to %s (%d bytes)", filepath, len(pdf_bytes))
    return filepath


# ---------------------------------------------------------------------------
# Resume exports
# ---------------------------------------------------------------------------

def export_resume_to_plain_text(resume: dict[str, Any]) -> str:
    """
    Contract: render a Resume dict to a plain-text string.

    Format:
      NAME
      Email: email | Phone: phone | Location: location
      LinkedIn: linkedin | GitHub: github

      SUMMARY
      -------
      <summary text>

      EXPERIENCE
      ----------
      Title | Company | Dates
      - Description point 1
      - Description point 2

      EDUCATION
      ---------
      Degree | School | Dates

      SKILLS
      ------
      skill1, skill2, skill3

      PROJECTS
      --------
      Project Name
      - Description
      - Tech: tech1, tech2

    Args:
        resume: a validated Resume dict with keys 'contact', 'summary',
                'experience', 'education', 'skills', 'projects'.

    Returns:
        A plain-text string.

    Raises:
        ValueError: if the resume is missing required fields.
    """
    contact = resume.get("contact", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    projects = resume.get("projects", [])

    if not contact.get("name"):
        raise ValueError("Resume must have a contact.name field")
    if not experience:
        raise ValueError("Resume must contain at least one experience entry")

    lines = []
    name = contact.get("name", "Unknown")
    lines.append(name.upper())
    lines.append("=" * len(name))
    lines.append("")

    # Contact info
    email = contact.get("email", "")
    phone = contact.get("phone", "")
    location = contact.get("location", "")
    linkedin = contact.get("linkedin", "")
    github = contact.get("github", "")

    contact_parts = [p for p in [email, phone, location, linkedin, github] if p]
    if contact_parts:
        lines.append(" | ".join(contact_parts))
        lines.append("")

    # Summary
    if summary:
        lines.append("SUMMARY")
        lines.append("-" * 7)
        lines.append(summary)
        lines.append("")

    # Experience
    if experience:
        lines.append("EXPERIENCE")
        lines.append("-" * 10)
        for exp in experience:
            title = exp.get("title", "Title")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "Dates")
            description = exp.get("description", "")
            lines.append(f"{title} | {company} | {dates}")
            # Bullet points from description
            for line in description.split("\n"):
                line = line.strip()
                if line:
                    lines.append(f"  - {line}")
            lines.append("")

    # Education
    if education:
        lines.append("EDUCATION")
        lines.append("-" * 9)
        for edu in education:
            degree = edu.get("degree", "Degree")
            school = edu.get("school", "School")
            dates = edu.get("dates", "Dates")
            lines.append(f"{degree} | {school} | {dates}")
        lines.append("")

    # Skills
    if skills:
        lines.append("SKILLS")
        lines.append("-" * 6)
        lines.append(", ".join(skills))
        lines.append("")

    # Projects
    if projects:
        lines.append("PROJECTS")
        lines.append("-" * 8)
        for proj in projects:
            proj_name = proj.get("name", "Project")
            desc = proj.get("description", "")
            tech = proj.get("tech", [])
            lines.append(proj_name)
            if desc:
                lines.append(f"  {desc}")
            if tech:
                lines.append(f"  Tech: {', '.join(tech)}")
            lines.append("")

    return "\n".join(lines).rstrip()


def export_resume_to_pdf(resume: dict[str, Any]) -> bytes:
    """
    Contract: render a Resume dict to a PDF document (bytes).
    Uses fpdf2 to generate a professional resume PDF.

    Args:
        resume: a validated Resume dict.

    Returns:
        PDF bytes ready to write to disk or serve.

    Raises:
        ValueError: if the resume is missing required fields.
        ImportError: if fpdf2 is not installed.
    """
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ImportError(
            "pdfexport requires fpdf2: pip install fpdf2"
        ) from exc

    contact = resume.get("contact", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    projects = resume.get("projects", [])

    if not contact.get("name"):
        raise ValueError("Resume must have a contact.name field")
    if not experience:
        raise ValueError("Resume must contain at least one experience entry")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)
    w = pdf.w - pdf.l_margin - pdf.r_margin  # usable width

    # Name
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(w, 12, contact.get("name", "Name"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Contact info
    pdf.set_font("Helvetica", "", 10)
    contact_parts = []
    for key in ["email", "phone", "location", "linkedin", "github"]:
        val = contact.get(key, "")
        if val:
            contact_parts.append(val)
    if contact_parts:
        pdf.cell(w, 6, "  |  ".join(contact_parts), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(6)

    # Summary
    if summary:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(w, 8, "SUMMARY", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(w, 6, summary, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # Experience
    if experience:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(w, 8, "EXPERIENCE", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        for exp in experience:
            title = exp.get("title", "Title")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "Dates")
            description = exp.get("description", "")

            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(w, 6, f"{title} | {company} | {dates}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            for line in description.split("\n"):
                line = line.strip()
                if line:
                    pdf.multi_cell(w, 5, f"- {line}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

    # Education
    if education:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(w, 8, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        for edu in education:
            degree = edu.get("degree", "Degree")
            school = edu.get("school", "School")
            dates = edu.get("dates", "Dates")
            pdf.cell(w, 6, f"{degree} | {school} | {dates}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # Skills
    if skills:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(w, 8, "SKILLS", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(w, 6, ", ".join(skills), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # Projects
    if projects:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(w, 8, "PROJECTS", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        for proj in projects:
            proj_name = proj.get("name", "Project")
            desc = proj.get("description", "")
            tech = proj.get("tech", [])
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(w, 6, proj_name, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            if desc:
                pdf.multi_cell(w, 5, desc, new_x="LMARGIN", new_y="NEXT")
            if tech:
                pdf.cell(w, 5, f"Tech: {', '.join(tech)}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

    # Serialize to bytes
    buf = BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def export_resume_to_docx(resume: dict[str, Any]) -> bytes:
    """
    Contract: render a Resume dict to a DOCX document (bytes).
    Uses python-docx to generate a professional resume document.

    Args:
        resume: a validated Resume dict.

    Returns:
        DOCX bytes ready to write to disk.

    Raises:
        ValueError: if the resume is missing required fields.
        ImportError: if python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError(
            "docx export requires python-docx: pip install python-docx"
        ) from exc

    contact = resume.get("contact", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    projects = resume.get("projects", [])

    if not contact.get("name"):
        raise ValueError("Resume must have a contact.name field")
    if not experience:
        raise ValueError("Resume must contain at least one experience entry")

    doc = Document()

    # Name
    name = contact.get("name", "Name")
    title_para = doc.add_heading(name, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact_parts = []
    for key in ["email", "phone", "location", "linkedin", "github"]:
        val = contact.get(key, "")
        if val:
            contact_parts.append(val)
    if contact_parts:
        contact_para = doc.add_paragraph("  |  ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    if summary:
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(summary)

    # Experience
    if experience:
        doc.add_heading("EXPERIENCE", level=1)
        for exp in experience:
            title = exp.get("title", "Title")
            company = exp.get("company", "Company")
            dates = exp.get("dates", "Dates")
            description = exp.get("description", "")

            exp_para = doc.add_paragraph()
            exp_para.add_run(f"{title} | {company} | {dates}").bold = True

            for line in description.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(f"- {line}", style="List Bullet")

    # Education
    if education:
        doc.add_heading("EDUCATION", level=1)
        for edu in education:
            degree = edu.get("degree", "Degree")
            school = edu.get("school", "School")
            dates = edu.get("dates", "Dates")
            doc.add_paragraph(f"{degree} | {school} | {dates}")

    # Skills
    if skills:
        doc.add_heading("SKILLS", level=1)
        doc.add_paragraph(", ".join(skills))

    # Projects
    if projects:
        doc.add_heading("PROJECTS", level=1)
        for proj in projects:
            proj_name = proj.get("name", "Project")
            desc = proj.get("description", "")
            tech = proj.get("tech", [])

            proj_para = doc.add_paragraph()
            proj_para.add_run(proj_name).bold = True

            if desc:
                doc.add_paragraph(desc)
            if tech:
                doc.add_paragraph(f"Tech: {', '.join(tech)}")

    # Serialize to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_resume_to_text_file(resume: dict[str, Any], filepath: str) -> str:
    """
    Contract: write resume plain-text export to a file.

    Args:
        resume: a validated Resume dict.
        filepath: path to write the .txt file.

    Returns:
        The filepath that was written.
    """
    text = export_resume_to_plain_text(resume)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)
    logger.info("Wrote resume plain-text export to %s", filepath)
    return filepath


def export_resume_to_pdf_file(resume: dict[str, Any], filepath: str) -> str:
    """
    Contract: write resume PDF export to a file.

    Args:
        resume: a validated Resume dict.
        filepath: path to write the .pdf file.

    Returns:
        The filepath that was written.
    """
    pdf_bytes = export_resume_to_pdf(resume)
    with open(filepath, "wb") as f:
        f.write(pdf_bytes)
    logger.info("Wrote resume PDF export to %s (%d bytes)", filepath, len(pdf_bytes))
    return filepath


def export_resume_to_docx_file(resume: dict[str, Any], filepath: str) -> str:
    """
    Contract: write resume DOCX export to a file.

    Args:
        resume: a validated Resume dict.
        filepath: path to write the .docx file.

    Returns:
        The filepath that was written.
    """
    docx_bytes = export_resume_to_docx(resume)
    with open(filepath, "wb") as f:
        f.write(docx_bytes)
    logger.info("Wrote resume DOCX export to %s (%d bytes)", filepath, len(docx_bytes))
    return filepath


# ---------------------------------------------------------------------------
# Audio exports
# ---------------------------------------------------------------------------

def export_audio_narrate(
    text: str,
    *,
    format: str = "wav",
    include_mp3: bool = True,
    output_path: str | None = None,
) -> dict[str, Any]:
    """
    Contract: narrate text to audio via audio-engine.

    Returns a dict with keys 'wav_bytes', 'mp3_bytes', 'sample_rate',
    'backend_used', 'voice_used'.
    """
    from engines.audio_engine.narration import narrate
    result = narrate(text, include_mp3=include_mp3, output_path=output_path)
    return {
        "wav_bytes": result.wav_bytes,
        "mp3_bytes": result.mp3_bytes,
        "sample_rate": result.sample_rate,
        "backend_used": result.backend_used,
        "voice_used": result.voice_used,
    }


def export_audio_podcast(
    topic: str,
    *,
    format: str = "wav",
    include_mp3: bool = True,
    output_path: str | None = None,
    num_segments: int = 5,
    duration_minutes: int = 10,
    host_name: str = "Host",
) -> dict[str, Any]:
    """
    Contract: generate a podcast script from a topic and render to audio.

    Returns a dict with keys 'wav_bytes', 'mp3_bytes', 'duration_seconds',
    'total_segments', 'backend_used'.
    """
    from engines.audio_engine.podcast_audio import render_podcast_to_audio
    from engines.audio_engine.podcast_script import generate_podcast_script
    script = generate_podcast_script(
        topic=topic,
        num_segments=num_segments,
        duration_minutes=duration_minutes,
        host_name=host_name,
    )
    audio = render_podcast_to_audio(
        script,
        include_mp3=include_mp3,
        output_path=output_path,
        speed=0.9,
    )
    return {
        "wav_bytes": audio.wav_bytes,
        "mp3_bytes": audio.mp3_bytes,
        "duration_seconds": audio.duration_seconds,
        "total_segments": audio.total_segments,
        "backend_used": audio.backend_used,
    }


def export_audio_bilingual(
    topic: str,
    target_language: str,
    known_language: str,
    *,
    format: str = "wav",
    include_mp3: bool = True,
    output_path: str | None = None,
    num_segments: int = 10,
) -> dict[str, Any]:
    """
    Contract: generate a bilingual pair and render to audio.

    Returns a dict with keys 'wav_bytes', 'mp3_bytes', 'duration_seconds',
    'total_segments', 'backend_used'.
    """
    from engines.language_lab.bilingual import (
        generate_bilingual_pair,
        render_bilingual_audio,
    )
    pair = generate_bilingual_pair(
        topic, target_language, known_language, num_segments=num_segments
    )
    audio = render_bilingual_audio(
        pair, include_mp3=include_mp3, output_path=output_path, speed=0.9
    )
    return {
        "wav_bytes": audio.wav_bytes,
        "mp3_bytes": audio.mp3_bytes,
        "duration_seconds": audio.duration_seconds,
        "total_segments": audio.total_segments,
        "backend_used": audio.backend_used,
    }


def export_audio_immersion(
    topic: str,
    target_language: str,
    *,
    format: str = "wav",
    include_mp3: bool = True,
    output_path: str | None = None,
    num_segments: int = 5,
    duration_minutes: int = 10,
) -> dict[str, Any]:
    """
    Contract: generate an immersion podcast and render to audio.

    Returns a dict with keys 'wav_bytes', 'mp3_bytes', 'duration_seconds',
    'total_segments', 'backend_used'.
    """
    from engines.language_lab.immersion import generate_immersion_podcast
    result = generate_immersion_podcast(
        topic,
        target_language,
        num_segments=num_segments,
        duration_minutes=duration_minutes,
        include_mp3=include_mp3,
        output_path=output_path,
        speed=0.9,
    )
    return {
        "wav_bytes": result.wav_bytes,
        "mp3_bytes": result.mp3_bytes,
        "duration_seconds": result.duration_seconds,
        "total_segments": result.total_segments,
        "backend_used": "PIPER",
    }


# ---------------------------------------------------------------------------
# Unified export dispatcher
# ---------------------------------------------------------------------------

def export(content: dict[str, Any], format: str = "text", filepath: str | None = None) -> str | bytes:
    """
    Contract: unified export function that detects content type and
    dispatches to the appropriate exporter.

    Args:
        content: either a Journey dict, a Resume dict, or an audio-type dict.
        format: output format — "text", "pdf", "docx", "wav", or "mp3".
        filepath: optional path to write to file.

    Returns:
        For in-memory: the exported content (str or bytes).
        For file: the filepath that was written.

    Raises:
        ValueError: if content type cannot be detected.
        KeyError: if format is not supported.
    """
    content_type = _detect_type(content)

    if content_type == "journey":
        if format == "text":
            result = export_to_plain_text(content)
        elif format == "pdf":
            result = export_to_pdf(content)
        else:
            raise KeyError(f"Journey export does not support format '{format}'")
    elif content_type == "resume":
        if format == "text":
            result = export_resume_to_plain_text(content)
        elif format == "pdf":
            result = export_resume_to_pdf(content)
        elif format == "docx":
            result = export_resume_to_docx(content)
        else:
            raise KeyError(f"Resume export does not support format '{format}'")
    elif content_type in ("narration", "podcast", "bilingual", "immersion"):
        format_lower = format.lower()
        if format_lower == "wav":
            result = export_audio(content, include_mp3=False)
        elif format_lower == "mp3":
            result = export_audio(content, include_mp3=True)
        else:
            raise KeyError(f"Audio export does not support format '{format}'")
    else:
        raise ValueError(f"Cannot detect content type. Keys: {list(content.keys())}")

    # Write to file if filepath provided
    if filepath:
        if content_type == "journey":
            if format == "text":
                export_to_text_file(content, filepath)
            elif format == "pdf":
                export_to_pdf_file(content, filepath)
        elif content_type == "resume":
            if format == "text":
                export_resume_to_text_file(content, filepath)
            elif format == "pdf":
                export_resume_to_pdf_file(content, filepath)
            elif format == "docx":
                export_resume_to_docx_file(content, filepath)
        elif content_type in ("narration", "podcast", "bilingual", "immersion"):
            export_audio(content, include_mp3=(format.lower() == "mp3"), output_path=filepath)
            return filepath
        return filepath

    return result


def export_audio(content: dict[str, Any], include_mp3: bool = True, output_path: str | None = None) -> dict[str, Any]:
    """
    Contract: route audio content (narration, podcast, bilingual, immersion)
    to the appropriate audio-engine function.

    Args:
        content: a dict detected as an audio type.
        include_mp3: whether to include MP3 bytes.
        output_path: optional directory to write audio files.

    Returns:
        Dict with audio result keys.

    Raises:
        ValueError: if content type cannot be determined.
    """
    content_type = _detect_type(content)
    if content_type not in ("narration", "podcast", "bilingual", "immersion"):
        raise ValueError(f"Content is not an audio type: {content_type}")

    if content_type == "narration":
        return export_audio_narrate(
            text=content["text"],
            include_mp3=include_mp3,
            output_path=output_path,
        )
    elif content_type == "podcast":
        return export_audio_podcast(
            topic=content["topic"],
            host_name=content.get("host_name", "Host"),
            num_segments=content.get("num_segments", 5),
            duration_minutes=content.get("duration_minutes", 10),
            include_mp3=include_mp3,
            output_path=output_path,
        )
    elif content_type == "bilingual":
        return export_audio_bilingual(
            topic=content["topic"],
            target_language=content["target_language"],
            known_language=content["known_language"],
            num_segments=content.get("num_segments", 10),
            include_mp3=include_mp3,
            output_path=output_path,
        )
    elif content_type == "immersion":
        return export_audio_immersion(
            topic=content["topic"],
            target_language=content["target_language"],
            num_segments=content.get("num_segments", 5),
            duration_minutes=content.get("duration_minutes", 10),
            include_mp3=include_mp3,
            output_path=output_path,
        )
    raise ValueError(f"Unknown audio content type: {content_type}")


# ---------------------------------------------------------------------------
# Re-export
# ---------------------------------------------------------------------------

__all__ = [
    # Journey exports
    "export_to_plain_text",
    "export_to_pdf",
    "export_to_text_file",
    "export_to_pdf_file",
    # Resume exports
    "export_resume_to_plain_text",
    "export_resume_to_pdf",
    "export_resume_to_docx",
    "export_resume_to_text_file",
    "export_resume_to_pdf_file",
    "export_resume_to_docx_file",
    # Audio exports
    "export_audio_narrate",
    "export_audio_podcast",
    "export_audio_bilingual",
    "export_audio_immersion",
    "export_audio",
    # Unified dispatcher
    "export",
    "_detect_type",
]
